import fitz
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extrair_data_assinatura_magistrado(pdf_path, paginas_alvo=[1, 2, 3, 4, 5]):
    """
    Extrai a data de assinatura digital (e-SAJ/TJMS) das margens ou metadados do PDF.
    """
    try:
        doc = fitz.open(pdf_path)
        data_assinatura = "não localizado"
        
        for p_num in paginas_alvo:
            if p_num > len(doc): break
            page = doc[p_num - 1]
            
            # O e-SAJ coloca assinaturas nas margens (texto rotacionado)
            blocos = page.get_text("blocks")
            for b in blocos:
                texto_bloco = b[4].upper()
                # Busca padrão de data (DD/MM/AAAA) próximo a "ASSINADO DIGITALMENTE" ou "TJMS"
                if "ASSINADO DIGITALMENTE" in texto_bloco or "TJMS" in texto_bloco:
                    match = re.search(r"(\d{2}/\d{2}/\d{4})", texto_bloco)
                    if match:
                        data_assinatura = f"{match.group(1)}"
                        doc.close()
                        return data_assinatura, p_num
        doc.close()
    except Exception as e:
        print(f"Erro ao extrair data de {pdf_path}: {e}")
        
    return "não localizado", None

def gerar_relatorio_word(dados_lista, nome_arquivo="Relatorio_Auditoria_Processual.docx"):
    """
    Gera um relatório Word consolidado com os dados da auditoria.
    """
    doc = Document()

    # Cabeçalho Estilizado
    titulo = doc.add_heading('RELATÓRIO DE AUDITORIA PROCESSUAL', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. RESUMO DOS ATOS PROCESSUAIS', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Folha'
    hdr_cells[1].text = 'Data Assinatura'
    hdr_cells[2].text = 'Tipo Documento'
    hdr_cells[3].text = 'Personagem/Réu'

    for dado in dados_lista:
        row_cells = table.add_row().cells
        row_cells[0].text = str(dado.get('folha', 'N/A'))
        row_cells[1].text = str(dado.get('data_assinatura', 'N/A'))
        row_cells[2].text = str(dado.get('tipo_conteudo', 'N/A'))
        row_cells[3].text = str(dado.get('personagem', 'N/A'))

    # Seção Detalhada (se houver narrativa)
    if any(d.get('narrativa_fato') for d in dados_lista):
        doc.add_page_break()
        doc.add_heading('2. DETALHAMENTO DOS FATOS', level=1)
        for dado in dados_lista:
            if dado.get('narrativa_fato'):
                doc.add_heading(f"Documento fls. {dado['folha']}", level=2)
                doc.add_paragraph(f"Réu/Personagem: {dado['personagem']}")
                doc.add_paragraph(dado['narrativa_fato'])
                if dado.get('atos_descritos_verbatim'):
                    doc.add_heading('Trecho Relevante (Verbatim)', level=3)
                    doc.add_paragraph(f'"{dado["atos_descritos_verbatim"]}"', style='Intense Quote')

    doc.save(nome_arquivo)
    print(f"✅ Relatório gerado com sucesso: {nome_arquivo}")
    return nome_arquivo
